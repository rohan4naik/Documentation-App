from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from app.services.supabase_service import SupabaseService
from app.services.ai_service import AIService
import traceback

router = APIRouter()
supabase_service = SupabaseService()
ai_service = AIService()

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    workspace_id: str = Form("default"),
    user_id: str = Form("anonymous")
):
    try:
        # 1. Read document content
        contents = await file.read()
        
        # Try to decode as text; if it fails (e.g. PDF), use pypdf if it's a PDF
        text_content = ""
        if file.filename.lower().endswith(".pdf"):
            try:
                import io
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(contents))
                extracted_text = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text.append(text)
                text_content = "\n".join(extracted_text)
                if not text_content.strip():
                    text_content = f"[Empty PDF file: {file.filename}] No text content could be extracted."
            except Exception as e:
                text_content = f"[Error reading PDF: {file.filename}] {str(e)}"
        else:
            try:
                text_content = contents.decode("utf-8")
            except UnicodeDecodeError:
                text_content = f"[Binary file: {file.filename}] Unable to extract text content."
        
        word_count = len(text_content.split())

        # 2. Analyze with Gemini
        ai_analysis = await ai_service.analyze_document(text_content)

        # 3. Insert Document row in Supabase database
        tags = ai_analysis.get("tags", [])
        if file.filename.lower().endswith(".pdf") and "pdf" not in tags:
            tags.append("pdf")

        doc_data = {
            "title": file.filename.rsplit('.', 1)[0] if '.' in file.filename else file.filename,
            "content": text_content[:10000],
            "excerpt": ai_analysis.get("summary", ""),
            "category": ai_analysis.get("category", "Uncategorized"),
            "tags": tags,
            "word_count": word_count,
        }

        # Add workspace_id and owner_id if they are valid UUIDs
        import uuid
        try:
            uuid.UUID(workspace_id)
            doc_data["workspace_id"] = workspace_id
        except ValueError:
            pass

        try:
            uuid.UUID(user_id)
            doc_data["owner_id"] = user_id
        except ValueError:
            pass

        try:
            new_doc = supabase_service.create_document(doc_data)
        except Exception as db_err:
            # Fallback if owner_id is missing from auth.users (e.g. cached old session or trigger delay)
            if "owner_id" in doc_data and ("23503" in str(db_err) or "owner_id" in str(db_err)):
                print(f"Warning: owner_id {doc_data['owner_id']} is not present in auth.users. Uploading document with owner_id = None.")
                del doc_data["owner_id"]
                new_doc = supabase_service.create_document(doc_data)
            else:
                raise db_err

        # 4. Upload original file to storage bucket
        try:
            is_pdf = file.filename.lower().endswith(".pdf") or contents.startswith(b"%PDF-")
            file_extension = ".pdf" if is_pdf else ""
            file_path = f"{new_doc['id']}{file_extension}"
            
            if is_pdf:
                content_type = "application/pdf"
            else:
                content_type = file.content_type or "application/octet-stream"
                
            supabase_service.upload_file(contents, file_path, content_type)
        except Exception as upload_err:
            print(f"Warning: Failed to upload file to storage bucket: {upload_err}")
            traceback.print_exc()

        return new_doc
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/convert")
async def convert_document(
    file: UploadFile = File(...),
    target_format: str = Form(...) # 'pdf', 'txt', 'md', 'docx', 'jpg'
):
    try:
        # Read the file content
        contents = await file.read()
        filename = file.filename
        
        # Detect input format
        input_ext = filename.split('.')[-1].lower() if '.' in filename else ''
        
        # 1. Parse text from input file
        text_content = ""
        if input_ext == "pdf" or contents.startswith(b"%PDF-"):
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            extracted_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text.append(text)
            text_content = "\n".join(extracted_text)
        else:
            try:
                text_content = contents.decode("utf-8")
            except UnicodeDecodeError:
                text_content = contents.decode("utf-8", errors="ignore")

        # 2. Convert text to target format
        out_filename = filename.rsplit('.', 1)[0] if '.' in filename else filename
        out_filename = f"{out_filename}.{target_format}"
        
        out_media_type = "application/octet-stream"
        
        if target_format == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            
            clean_text = text_content.encode("latin-1", errors="replace").decode("latin-1")
            
            for line in clean_text.splitlines():
                pdf.multi_cell(0, 10, txt=line)
            
            out_bytes = bytes(pdf.output())
            out_media_type = "application/pdf"
            
        elif target_format == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument()
            doc.add_heading(filename.rsplit('.', 1)[0], level=1)
            for line in text_content.splitlines():
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph("")
            doc_io = io.BytesIO()
            doc.save(doc_io)
            out_bytes = doc_io.getvalue()
            out_media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        elif target_format == "jpg":
            if not (input_ext == "pdf" or contents.startswith(b"%PDF-")):
                raise HTTPException(status_code=400, detail="JPG conversion is only supported for PDF input files.")
            
            import fitz  # PyMuPDF
            doc = fitz.open(stream=contents, filetype="pdf")
            num_pages = len(doc)
            
            if num_pages == 0:
                raise HTTPException(status_code=400, detail="Empty PDF file.")
            elif num_pages == 1:
                page = doc.load_page(0)
                pix = page.get_pixmap()
                out_bytes = bytes(pix.tobytes("jpg"))
                out_media_type = "image/jpeg"
            else:
                import zipfile
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                    for i in range(num_pages):
                        page = doc.load_page(i)
                        pix = page.get_pixmap()
                        img_data = bytes(pix.tobytes("jpg"))
                        zip_file.writestr(f"page_{i+1}.jpg", img_data)
                
                out_bytes = zip_buffer.getvalue()
                out_media_type = "application/zip"
                out_filename = f"{out_filename.rsplit('.', 1)[0]}_images.zip"
            
        elif target_format == "md":
            md_content = f"# {filename.rsplit('.', 1)[0]}\n\n{text_content}"
            out_bytes = md_content.encode("utf-8")
            out_media_type = "text/markdown"
            
        elif target_format == "txt":
            out_bytes = text_content.encode("utf-8")
            out_media_type = "text/plain"
            
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported target format: {target_format}")

        import base64
        base64_str = base64.b64encode(out_bytes).decode("utf-8")
        return {
            "filename": out_filename,
            "media_type": out_media_type,
            "base64": base64_str
        }
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
